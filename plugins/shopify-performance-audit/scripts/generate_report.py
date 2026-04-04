#!/usr/bin/env python3
"""
Shopify Performance Audit Report Generator
Generates a .docx performance audit report for any Shopify website.

Usage:
    python3 generate_report.py \
        --homepage "https://www.example.com/" \
        --pdp "https://www.example.com/products/some-product" \
        --plp "https://www.example.com/collections/some-collection" \
        --brand "Brand Name" \
        --logo "/path/to/logo.png" (optional) \
        --output "/path/to/output.docx" (optional)
"""

import argparse
import json
import os
import sys
import time
import urllib.request
import urllib.parse
import urllib.error
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--user", "-q"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml


# ─── Ideal Benchmark Ranges ───────────────────────────────────────────
BENCHMARKS = {
    "cls": {"mobile": 0.1, "desktop": 0.1},
    "tbt": {"mobile": 800, "desktop": 600},  # ms
    "fcp": {"mobile": 2.0, "desktop": 1.8},  # sec
    "lcp": {"mobile": 3.0, "desktop": 2.5},  # sec
    "si":  {"mobile": 4.0, "desktop": 3.0},  # sec
}

SCOPE_THRESHOLDS = {
    "cls": {
        "mobile":  [(0.1, "Already good"), (0.25, "Minor scope of improvements"), (0.5, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
        "desktop": [(0.1, "Already good"), (0.25, "Minor scope of improvements"), (0.5, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
    },
    "tbt": {
        "mobile":  [(800, "Already good"), (1500, "Minor scope of improvements"), (3000, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
        "desktop": [(600, "Already good"), (1000, "Minor scope of improvements"), (2000, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
    },
    "fcp": {
        "mobile":  [(2, "Already good"), (4, "Minor scope of improvements"), (6, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
        "desktop": [(1.8, "Already good"), (3, "Minor scope of improvements"), (5, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
    },
    "lcp": {
        "mobile":  [(3, "Already good"), (6, "Minor scope of improvements"), (12, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
        "desktop": [(2.5, "Already good"), (4, "Minor scope of improvements"), (8, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
    },
    "si": {
        "mobile":  [(4, "Already good"), (6, "Minor scope of improvements"), (10, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
        "desktop": [(3, "Already good"), (5, "Minor scope of improvements"), (8, "Scope of improvements"), (float('inf'), "Major scope of improvements")],
    },
}


def get_scope(metric, device, value):
    """Determine the scope of improvement for a metric value."""
    thresholds = SCOPE_THRESHOLDS[metric][device]
    for threshold, label in thresholds:
        if value <= threshold:
            return label
    return "Major scope of improvements"


# Default API key for PageSpeed Insights (avoids shared quota / 429 rate limits)
GOOGLE_API_KEY = "AIzaSyBIjNlne770dB5IiTzHOGM1A2NxNBd8674"


def make_cache_busted_url(url, test_index):
    """Append a unique cache-busting query parameter to avoid duplicate/cached results."""
    bust_value = int(time.time()) + test_index
    separator = "&" if "?" in url else "?"
    return f"{url}{separator}_psi_bust={bust_value}"


def run_pagespeed_test(url, strategy="mobile", test_index=0):
    """Run a single PageSpeed Insights test and return results.
    Uses cache-busting to avoid cookie/CDN/API caching issues."""
    busted_url = make_cache_busted_url(url, test_index)
    api_url = (
        f"https://www.googleapis.com/pagespeedonline/v5/runPagespeed"
        f"?url={urllib.parse.quote(busted_url, safe='')}"
        f"&strategy={strategy}"
        f"&category=performance"
    )
    if GOOGLE_API_KEY:
        api_url += f"&key={GOOGLE_API_KEY}"

    last_error = "Unknown error"
    max_attempts = 4  # More retries for 429 rate limiting
    for attempt in range(max_attempts):
        try:
            req = urllib.request.Request(api_url, headers={
                "User-Agent": "ShopifyAuditBot/1.0",
                "Cache-Control": "no-cache",
                "Pragma": "no-cache",
            })
            with urllib.request.urlopen(req, timeout=180) as response:
                data = json.loads(response.read().decode())

            lhr = data.get("lighthouseResult", {})
            audits = lhr.get("audits", {})
            categories = lhr.get("categories", {})

            perf_score = categories.get("performance", {}).get("score", 0)
            if perf_score is not None:
                perf_score = round(perf_score * 100)
            else:
                perf_score = 0

            cls_val = audits.get("cumulative-layout-shift", {}).get("numericValue", 0)
            tbt_val = audits.get("total-blocking-time", {}).get("numericValue", 0)  # ms
            fcp_val = audits.get("first-contentful-paint", {}).get("numericValue", 0) / 1000  # ms -> sec
            lcp_val = audits.get("largest-contentful-paint", {}).get("numericValue", 0) / 1000  # ms -> sec
            si_val = audits.get("speed-index", {}).get("numericValue", 0) / 1000  # ms -> sec

            # Core Web Vitals assessment
            cwv_passed = True
            # Check CLS
            if cls_val > 0.25:
                cwv_passed = False
            # Check LCP
            if lcp_val > 4.0:
                cwv_passed = False
            # Check TBT (proxy for INP)
            if tbt_val > 600:
                cwv_passed = False

            # Build the shareable PageSpeed analysis link
            # Format: https://pagespeed.web.dev/analysis?url={url}&form_factor=mobile
            analysis_link = (
                f"https://pagespeed.web.dev/analysis"
                f"?url={urllib.parse.quote(url, safe='')}"
                f"&form_factor={strategy}"
            )

            return {
                "score": perf_score,
                "cls": round(cls_val, 3),
                "tbt": round(tbt_val),
                "fcp": round(fcp_val, 1),
                "lcp": round(lcp_val, 1),
                "si": round(si_val, 1),
                "cwv_passed": cwv_passed,
                "analysis_url": analysis_link,
                "success": True,
            }

        except urllib.error.HTTPError as e:
            last_error = str(e)
            print(f"  Attempt {attempt + 1}/{max_attempts} failed: {e}")
            if e.code == 429:
                # Rate limited — exponential backoff: 30s, 60s, 120s
                wait_time = 30 * (2 ** attempt)
                print(f"  Rate limited (429). Waiting {wait_time} seconds before retry...")
                time.sleep(wait_time)
            elif attempt < max_attempts - 1:
                print(f"  Retrying in 15 seconds...")
                time.sleep(15)
        except Exception as e:
            last_error = str(e)
            print(f"  Attempt {attempt + 1}/{max_attempts} failed: {e}")
            if attempt < max_attempts - 1:
                print(f"  Retrying in 15 seconds...")
                time.sleep(15)

    return {"success": False, "error": last_error}


def validate_unique_results(tests, strategy, page_name):
    """Check if all test scores are identical (likely cached). Returns True if unique enough."""
    if len(tests) < 2:
        return True
    scores = [t["score"] for t in tests]
    if len(set(scores)) == 1:
        print(f"  WARNING: All {len(tests)} {strategy} scores for {page_name} are identical ({scores[0]}). Likely cached results.")
        return False
    return True


def run_parallel_mobile_desktop(url, test_index):
    """Run mobile and desktop tests in parallel for the same URL.
    Returns (mobile_result, desktop_result) tuple.
    This halves the total time: 1 parallel call = both mobile + desktop."""
    with ThreadPoolExecutor(max_workers=2) as executor:
        future_mobile = executor.submit(run_pagespeed_test, url, "mobile", test_index)
        future_desktop = executor.submit(run_pagespeed_test, url, "desktop", test_index + 100)
        mobile_result = future_mobile.result()
        desktop_result = future_desktop.result()
    return mobile_result, desktop_result


def run_tests_for_page(url, page_name, num_tests=3):
    """Run multiple PageSpeed tests for a page and return averaged results.
    Runs mobile + desktop in PARALLEL per test iteration = only 3 API rounds per page.
    Total API calls: still 6 per page (3 mobile + 3 desktop) but executed as 3 parallel pairs,
    so wall-clock time is equivalent to ~3 sequential calls instead of 6."""
    results = {"mobile": [], "desktop": [], "links": []}

    for i in range(num_tests):
        print(f"  Running {page_name} test {i+1}/{num_tests} (mobile + desktop in parallel)...")
        mobile_result, desktop_result = run_parallel_mobile_desktop(url, test_index=i)

        if mobile_result.get("success"):
            results["mobile"].append(mobile_result)
            # Build proper shareable link for the report
            link = (
                f"https://pagespeed.web.dev/analysis"
                f"?url={urllib.parse.quote(url, safe='')}"
                f"&form_factor=mobile"
            )
            results["links"].append(link)
        else:
            print(f"  WARNING: Mobile test failed for {page_name} test {i+1}")

        if desktop_result.get("success"):
            results["desktop"].append(desktop_result)
        else:
            print(f"  WARNING: Desktop test failed for {page_name} test {i+1}")

        # Delay between test rounds to avoid rate limiting and get fresh results
        if i < num_tests - 1:
            time.sleep(15)

    # Validate results are unique — retry if all scores are identical (cached)
    for strategy in ["mobile", "desktop"]:
        if not validate_unique_results(results[strategy], strategy, page_name):
            print(f"  Re-running one {strategy} test for {page_name} after 30s cooldown...")
            time.sleep(30)
            retry_result = run_pagespeed_test(url, strategy, test_index=99)
            if retry_result.get("success"):
                if len(results[strategy]) >= 2:
                    results[strategy][1] = retry_result
                else:
                    results[strategy].append(retry_result)

    # Calculate averages
    averaged = {}
    for strategy in ["mobile", "desktop"]:
        tests = results[strategy]
        if tests:
            averaged[strategy] = {
                "score": round(sum(t["score"] for t in tests) / len(tests)),
                "cls": round(sum(t["cls"] for t in tests) / len(tests), 3),
                "tbt": round(sum(t["tbt"] for t in tests) / len(tests)),
                "fcp": round(sum(t["fcp"] for t in tests) / len(tests), 1),
                "lcp": round(sum(t["lcp"] for t in tests) / len(tests), 1),
                "si": round(sum(t["si"] for t in tests) / len(tests), 1),
                "cwv_passed": sum(1 for t in tests if t["cwv_passed"]) > len(tests) / 2,
            }
        else:
            averaged[strategy] = {
                "score": 0, "cls": 0, "tbt": 0, "fcp": 0, "lcp": 0, "si": 0, "cwv_passed": False
            }

    averaged["links"] = results["links"]
    return averaged


def calculate_expected_after(current_mobile_score, current_desktop_score, cwv_mobile, cwv_desktop):
    """Calculate expected after optimization values."""
    expected_mobile = f"Expected: {min(current_mobile_score + 10, 65)}-{min(current_mobile_score + 20, 70)}+"
    expected_desktop = f"Expected: {min(current_desktop_score + 5, 85)}-{min(current_desktop_score + 15, 90)}+"
    expected_cwv_mobile = "Passed" if cwv_mobile else "Expected: Passed"
    expected_cwv_desktop = "Passed" if cwv_desktop else "Expected: Passed"
    return expected_mobile, expected_desktop, expected_cwv_mobile, expected_cwv_desktop


# ─── Document Building Helpers ─────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, style=None, bold=False, font_size=11, alignment=None, space_after=None, font_name="Calibri"):
    """Add a styled paragraph to the document."""
    para = doc.add_paragraph()
    if style:
        para.style = style
    run = para.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    if alignment:
        para.alignment = alignment
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def add_performance_table(doc, avg_mobile, avg_desktop, cwv_mobile, cwv_desktop, exp_mobile, exp_desktop, exp_cwv_mobile, exp_cwv_desktop):
    """Add the before/after performance table."""
    table = doc.add_table(rows=3, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["", "Average Mobile", "Average Desktop", "Core Web Vitals Mobile", "Core Web Vitals Desktop"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    # Before row
    before_values = [
        "Before",
        str(avg_mobile),
        str(avg_desktop),
        "Passed" if cwv_mobile else "Failed",
        "Passed" if cwv_desktop else "Failed",
    ]
    for i, val in enumerate(before_values):
        cell = table.rows[1].cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    # After row
    after_values = [
        "After speed opt.",
        exp_mobile,
        exp_desktop,
        exp_cwv_mobile,
        exp_cwv_desktop,
    ]
    for i, val in enumerate(after_values):
        cell = table.rows[2].cells[i]
        cell.text = val
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacing
    return table


def add_metric_section(doc, metric_name, metric_abbr, mobile_val, desktop_val, mobile_unit, mobile_scope, desktop_scope, description_text=None, improvement_points=None, extra_note=None):
    """Add a metric improvement section."""
    doc.add_heading(f"Improve {metric_abbr} ({metric_name})", level=2)

    # Current values
    mobile_display = f"{mobile_val} {mobile_unit}" if mobile_unit else str(mobile_val)
    desktop_display = f"{desktop_val} {mobile_unit}" if mobile_unit else str(desktop_val)

    para = doc.add_paragraph()
    run = para.add_run(f"Current {metric_abbr} (Mobile): {mobile_display} ({mobile_scope})\nCurrent {metric_abbr} (Desktop): {desktop_display} ({desktop_scope})")
    run.font.size = Pt(11)

    if description_text:
        para = doc.add_paragraph()
        run = para.add_run(description_text)
        run.font.size = Pt(11)

    if improvement_points:
        doc.add_heading(f"To improve {metric_abbr}:", level=2)
        for point in improvement_points:
            para = doc.add_paragraph(point, style="List Bullet")
            for run in para.runs:
                run.font.size = Pt(11)

    if extra_note:
        para = doc.add_paragraph()
        run = para.add_run(extra_note)
        run.font.size = Pt(11)


def add_pdp_plp_section(doc, page_name, data, links):
    """Add PDP or PLP findings section (simplified, no detailed explanations)."""
    doc.add_heading(f"{page_name} Findings:", level=1)
    doc.add_heading("Current speed performance reports:", level=2)

    for link in links[:3]:
        para = doc.add_paragraph()
        run = para.add_run(link)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    # Table
    mobile = data["mobile"]
    desktop = data["desktop"]
    exp_m, exp_d, exp_cwv_m, exp_cwv_d = calculate_expected_after(
        mobile["score"], desktop["score"], mobile["cwv_passed"], desktop["cwv_passed"]
    )
    add_performance_table(
        doc, mobile["score"], desktop["score"],
        mobile["cwv_passed"], desktop["cwv_passed"],
        exp_m, exp_d, exp_cwv_m, exp_cwv_d
    )

    doc.add_heading("What Can Be Done to Improve Website Speed", level=2)

    # CLS
    doc.add_paragraph("Improve CLS (Cumulative Layout Shift)", style="List Bullet")
    cls_m_scope = get_scope("cls", "mobile", mobile["cls"])
    cls_d_scope = get_scope("cls", "desktop", desktop["cls"])
    para = doc.add_paragraph()
    run = para.add_run(f"Current CLS (Mobile): {mobile['cls']} ({cls_m_scope})\nCurrent CLS (Desktop): {desktop['cls']} ({cls_d_scope})")
    run.font.size = Pt(11)

    # TBT
    doc.add_heading("Improve TBT (Total Blocking Time)", level=2)
    tbt_m_scope = get_scope("tbt", "mobile", mobile["tbt"])
    tbt_d_scope = get_scope("tbt", "desktop", desktop["tbt"])
    para = doc.add_paragraph()
    run = para.add_run(f"Current TBT (Mobile): {mobile['tbt']} ms ({tbt_m_scope})\nCurrent TBT (Desktop): {desktop['tbt']} ms ({tbt_d_scope})")
    run.font.size = Pt(11)

    # FCP
    doc.add_heading("Improve FCP (First Contentful Paint)", level=2)
    fcp_m_scope = get_scope("fcp", "mobile", mobile["fcp"])
    fcp_d_scope = get_scope("fcp", "desktop", desktop["fcp"])
    para = doc.add_paragraph()
    run = para.add_run(f"Current FCP (Mobile): {mobile['fcp']} sec ({fcp_m_scope})\nCurrent FCP (Desktop): {desktop['fcp']} sec ({fcp_d_scope})")
    run.font.size = Pt(11)

    # LCP
    doc.add_heading("Improve LCP (Largest Contentful Paint)", level=2)
    lcp_m_scope = get_scope("lcp", "mobile", mobile["lcp"])
    lcp_d_scope = get_scope("lcp", "desktop", desktop["lcp"])
    para = doc.add_paragraph()
    run = para.add_run(f"Current LCP (Mobile): {mobile['lcp']} sec ({lcp_m_scope})\nCurrent LCP (Desktop): {desktop['lcp']} sec ({lcp_d_scope})")
    run.font.size = Pt(11)

    # SI
    doc.add_heading("Improve SI (Speed Index)", level=2)
    si_m_scope = get_scope("si", "mobile", mobile["si"])
    si_d_scope = get_scope("si", "desktop", desktop["si"])
    para = doc.add_paragraph()
    run = para.add_run(f"Current SI (Mobile): {mobile['si']} sec ({si_m_scope})\nCurrent SI (Desktop): {desktop['si']} sec ({si_d_scope})")
    run.font.size = Pt(11)


def add_footer(doc):
    """Add the standard footer to all sections."""
    footer_text = (
        "Devxconsultancy Private Limited\n"
        "https://www.devxlabs.ai/\n\n"
        "U-56, Sandhya Darshan Apartment,\n"
        "Surat, Gujarat, India - 395009\n\n"
        "GSTIN:24AAKCD0230A1ZJ | PAN:AAKCD0230A"
    )
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.text = ""
        run = para.add_run(footer_text)
        run.font.size = Pt(8)
        run.font.name = "Calibri"
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_logo_to_header(doc, logo_path):
    """Add brand logo to document header."""
    if logo_path and Path(logo_path).exists():
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            try:
                run = para.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except Exception as e:
                print(f"  WARNING: Could not add logo: {e}")


def generate_report(homepage_data, pdp_data, plp_data, brand_name, logo_path=None, output_path=None):
    """Generate the complete .docx performance audit report."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ─── Cover / Title ───
    add_styled_paragraph(doc, "Website Performance Optimization", bold=True, font_size=24, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_styled_paragraph(doc, f"Brand: {brand_name}", font_size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_styled_paragraph(doc, "Submitted by: devx labs", font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

    doc.add_page_break()

    # ─── Executive Summary ───
    doc.add_heading("Executive Summary", level=1)
    para = doc.add_paragraph()
    run = para.add_run(
        f"{brand_name}'s website performance has been analysed across Homepage, Product Detail Pages (PDP), "
        f"and Collection Pages (PLP) to evaluate loading behaviour, visual stability, and user experience across "
        f"devices. The audit highlights that while desktop performance remains relatively stable, mobile performance "
        f"presents opportunities for improvement, particularly across rendering speed, loading prioritisation, and "
        f"asset delivery."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run(
        "The purpose of this optimisation initiative is to enhance overall performance efficiency while "
        "maintaining the existing design structure and functionality. The recommendations focus on improving "
        "Core Web Vitals stability, reducing delays during page load, and ensuring a smoother browsing experience "
        "across key customer journeys."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_page_break()

    # ─── Homepage Findings ───
    doc.add_heading("Website Performance Audit Report", level=1)
    doc.add_paragraph("Homepage Findings", style="List Bullet")
    doc.add_heading("Current speed performance reports:", level=2)

    for link in homepage_data.get("links", [])[:4]:
        para = doc.add_paragraph()
        run = para.add_run(link)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    hp_mobile = homepage_data["mobile"]
    hp_desktop = homepage_data["desktop"]

    exp_m, exp_d, exp_cwv_m, exp_cwv_d = calculate_expected_after(
        hp_mobile["score"], hp_desktop["score"],
        hp_mobile["cwv_passed"], hp_desktop["cwv_passed"]
    )

    add_performance_table(
        doc, hp_mobile["score"], hp_desktop["score"],
        hp_mobile["cwv_passed"], hp_desktop["cwv_passed"],
        exp_m, exp_d, exp_cwv_m, exp_cwv_d
    )

    doc.add_heading("What Can Be Done to Improve Website Speed", level=2)

    # CLS
    doc.add_paragraph("Improve CLS (Cumulative Layout Shift)", style="List Bullet")
    cls_m_scope = get_scope("cls", "mobile", hp_mobile["cls"])
    cls_d_scope = get_scope("cls", "desktop", hp_desktop["cls"])
    para = doc.add_paragraph()
    run = para.add_run(
        "Our goal is to keep CLS close to 0 (~0) to improve load time, ensure visual stability on first view, "
        "and enhance the overall user experience."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run(
        f"Current CLS (Mobile): {hp_mobile['cls']} ({cls_m_scope})\n"
        f"Current CLS (Desktop): {hp_desktop['cls']} ({cls_d_scope})"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_heading("To improve CLS:", level=2)
    doc.add_paragraph(
        "We need to ensure that no elements move or change their size or position during the page load process.",
        style="List Bullet"
    )

    # TBT
    doc.add_heading("Improve TBT (Total Blocking Time)", level=2)
    tbt_m_scope = get_scope("tbt", "mobile", hp_mobile["tbt"])
    tbt_d_scope = get_scope("tbt", "desktop", hp_desktop["tbt"])
    para = doc.add_paragraph()
    run = para.add_run(
        f"Current TBT (Mobile): {hp_mobile['tbt']} ms ({tbt_m_scope})\n"
        f"Current TBT (Desktop): {hp_desktop['tbt']} ms ({tbt_d_scope})"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run(
        "We need to reduce this blocking time to below 1000 ms, ideally around 500 ms (in optimal cases)."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_heading("To improve TBT:", level=2)
    doc.add_paragraph(
        "This can be achieved by optimizing scripts, CSS, and managing third-party integrations that contribute to blocking the main thread.",
        style="List Bullet"
    )

    # FCP
    doc.add_heading("Improve FCP (First Contentful Paint)", level=2)
    fcp_m_scope = get_scope("fcp", "mobile", hp_mobile["fcp"])
    fcp_d_scope = get_scope("fcp", "desktop", hp_desktop["fcp"])
    para = doc.add_paragraph()
    run = para.add_run(
        f"Current FCP (Mobile): {hp_mobile['fcp']} sec ({fcp_m_scope})\n"
        f"Current FCP (Desktop): {hp_desktop['fcp']} sec ({fcp_d_scope})"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run(
        "FCP measures how long it takes for the first visible element to render on page load. "
        "It depends largely on code quality, DOM size, and critical or third-party scripts."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_heading("To improve FCP:", level=2)
    doc.add_paragraph("Optimize and clean up code in the theme.liquid file.", style="List Bullet")
    doc.add_paragraph("Reduce and manage critical and third-party scripts as much as possible.", style="List Bullet")

    # LCP
    doc.add_heading("Improve LCP (Largest Contentful Paint)", level=2)
    lcp_m_scope = get_scope("lcp", "mobile", hp_mobile["lcp"])
    lcp_d_scope = get_scope("lcp", "desktop", hp_desktop["lcp"])
    para = doc.add_paragraph()
    run = para.add_run(
        f"Current LCP (Mobile): {hp_mobile['lcp']} sec ({lcp_m_scope})\n"
        f"Current LCP (Desktop): {hp_desktop['lcp']} sec ({lcp_d_scope})"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run("LCP is one of the most impactful metrics for page load speed.")
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_heading("To improve LCP:", level=2)
    doc.add_paragraph("Ensure the LCP element remains consistent on every page load.", style="List Bullet")
    doc.add_paragraph("Set fetch priority = high for the LCP element.", style="List Bullet")
    doc.add_paragraph("Set fetch priority = low for all other non-critical assets.", style="List Bullet")
    doc.add_paragraph("Lazy load all images except the LCP element to avoid heavy asset loading on initial render.", style="List Bullet")

    # FCP-LCP gap note
    gap = round(hp_mobile["lcp"] - hp_mobile["fcp"], 1)
    if gap > 5:
        para = doc.add_paragraph()
        run = para.add_run(
            f"There is a major gap between FCP and LCP on mobile (around {gap} seconds, "
            f"i.e., {hp_mobile['lcp']} sec - {hp_mobile['fcp']} sec).\n"
            f"This delay is primarily caused by slow LCP element loading and render delays, "
            f"which we need to optimize."
        )
        run.font.size = Pt(11)
        para.style = "Body Text"

    # SI
    doc.add_heading("Improve SI (Speed Index)", level=2)
    si_m_scope = get_scope("si", "mobile", hp_mobile["si"])
    si_d_scope = get_scope("si", "desktop", hp_desktop["si"])
    para = doc.add_paragraph()
    run = para.add_run(
        f"Current SI (Mobile): {hp_mobile['si']} sec ({si_m_scope})\n"
        f"Current SI (Desktop): {hp_desktop['si']} sec ({si_d_scope})"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    para = doc.add_paragraph()
    run = para.add_run(
        "Speed Index depends on the performance of all the above metrics, "
        "but is also influenced by other factors:"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_paragraph("Optimize or remove long-running JavaScript functions and network tasks.", style="List Bullet")
    doc.add_paragraph("Rearrange scripts in network calls based on priority and dependency.", style="List Bullet")
    doc.add_paragraph("Reduce inter-dependencies among scripts to eliminate waiting time.", style="List Bullet")
    doc.add_paragraph("Follow best coding practices to keep solutions minimal, efficient, and easy to execute.", style="List Bullet")

    # Additional Optimization Points
    doc.add_heading("Additional Optimization Points", level=2)
    doc.add_paragraph("Remove unused or unnecessary apps from the store.", style="List Bullet")
    doc.add_paragraph("Clean up legacy code from previously used apps that are no longer active.", style="List Bullet")
    doc.add_paragraph("Convert images to next-generation formats (WEBP, PNG) to improve render speed.", style="List Bullet")
    doc.add_paragraph("Resize and serve images according to actual display dimensions.", style="List Bullet")
    para = doc.add_paragraph()
    run = para.add_run(
        "For example, if a product card displays an image at 300\u00d7300, avoid loading a 1000\u00d71000 version, "
        "as it increases CPU load and slows down the page."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_page_break()

    # ─── Summary / Action Plan ───
    doc.add_heading("Summary / Action Plan", level=1)

    doc.add_heading("Cumulative Layout Shift (CLS)", level=2)
    doc.add_paragraph("Set fixed width & height for images, videos, and banners.", style="List Bullet")
    doc.add_paragraph("Predefine container sizes for all dynamic elements (reviews, icons, etc.).", style="List Bullet")
    doc.add_paragraph("Avoid inserting new elements above existing content after load.", style="List Bullet")

    doc.add_heading("Total Blocking Time (TBT)", level=2)
    doc.add_paragraph("Defer or async non-critical JavaScript files.", style="List Bullet")
    doc.add_paragraph("Minimize third-party scripts (tracking, analytics, chat widgets).", style="List Bullet")
    doc.add_paragraph("Move non-essential scripts to load on user interaction (scroll, click, etc.).", style="List Bullet")
    doc.add_paragraph("Split large JS bundles; optimize Shopify theme.js and custom scripts.", style="List Bullet")
    doc.add_paragraph("Remove unused CSS and JS from the theme.", style="List Bullet")

    doc.add_heading("First Contentful Paint (FCP)", level=2)
    doc.add_paragraph("Optimize theme.liquid by cleaning up unnecessary code and inline scripts.", style="List Bullet")
    doc.add_paragraph("Reduce critical JS/CSS and prioritize above-the-fold content.", style="List Bullet")
    doc.add_paragraph("Use <picture> tags and responsive images for mobile and desktop.", style="List Bullet")
    doc.add_paragraph("Limit the number of third-party scripts running on page load.", style="List Bullet")

    doc.add_heading("Largest Contentful Paint (LCP)", level=2)
    doc.add_paragraph("Identify and prioritize the main LCP element (usually banner image or hero text).", style="List Bullet")
    doc.add_paragraph('Set fetchpriority="high" for LCP element; fetchpriority="low" for others.', style="List Bullet")
    doc.add_paragraph("Lazy load all images except the LCP one.", style="List Bullet")
    doc.add_paragraph("Compress and serve LCP images in WEBP format.", style="List Bullet")
    doc.add_paragraph("Avoid animations or opacity transitions on the LCP element.", style="List Bullet")
    doc.add_paragraph("Ensure the LCP element appears within the first viewport and is consistent on reload.", style="List Bullet")

    doc.add_heading("Speed Index (SI)", level=2)
    doc.add_paragraph("Optimize overall code execution order (critical first).", style="List Bullet")
    doc.add_paragraph("Reorder scripts based on dependency and importance.", style="List Bullet")
    doc.add_paragraph("Remove long-running JS tasks and optimize loops or event listeners.", style="List Bullet")
    doc.add_paragraph("Keep DOM structure lightweight.", style="List Bullet")
    doc.add_paragraph("Follow clean, modular code practices.", style="List Bullet")

    doc.add_heading("Additional Optimizations", level=2)
    doc.add_paragraph("Remove unused or disabled apps from Shopify.", style="List Bullet")
    doc.add_paragraph("Clean up old app code snippets left in theme.liquid or assets.", style="List Bullet")
    doc.add_paragraph("Convert images to next-gen formats (WEBP/PNG).", style="List Bullet")
    doc.add_paragraph("Serve images in exact required sizes (no oversized assets).", style="List Bullet")

    doc.add_page_break()

    # ─── PDP Findings ───
    add_pdp_plp_section(doc, "PDP (Product Detail Page)", pdp_data, pdp_data.get("links", []))

    doc.add_page_break()

    # ─── PLP Findings ───
    add_pdp_plp_section(doc, "PLP (Product Listing Page)", plp_data, plp_data.get("links", []))

    doc.add_page_break()

    # ─── Common Performance Improvements ───
    doc.add_heading("Common Performance Improvements", level=1)
    para = doc.add_paragraph()
    run = para.add_run(
        "To improve your website speed and user experience, we've identified a few key areas "
        "that have the biggest impact on performance"
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_heading("Total Blocking Time (TBT)", level=2)
    doc.add_paragraph("Optimizing JS files: swiper and animations.", style="List Bullet")
    doc.add_paragraph("Cart and bundle-related scripts (including third-party apps)", style="List Bullet")
    doc.add_paragraph("Ensuring scripts load only when needed and do not block the main page.", style="List Bullet")

    doc.add_heading("First Contentful Paint (FCP)", level=2)
    doc.add_paragraph("Third-party apps or plugins: Marketing and tracking scripts", style="List Bullet")
    doc.add_paragraph("Improving image loading and delivery.", style="List Bullet")
    doc.add_paragraph("The first visible image loads immediately when the page opens.", style="List Bullet")
    doc.add_paragraph("Reviewing framework usage (e.g., React) to minimize its impact on initial load time.", style="List Bullet")

    doc.add_heading("Largest Contentful Paint (LCP)", level=2)
    doc.add_paragraph("Render-blocking resources.", style="List Bullet")
    doc.add_paragraph("Fixing image loading priorities so key visuals load first.", style="List Bullet")
    doc.add_paragraph("Optimizing image sizes and delivery for different devices.", style="List Bullet")

    para = doc.add_paragraph()
    run = para.add_run(
        "Note: Core Web Vitals data on Google updates approximately 28 days after the site goes live, "
        "so we'll need to wait for that period after go-live to see the actual impact from real user data."
    )
    run.font.size = Pt(11)
    run.italic = True
    para.style = "Body Text"

    doc.add_page_break()

    # ─── Ideal Performance Benchmark Range ───
    doc.add_heading("Ideal Performance Benchmark Range", level=1)
    bench_table = doc.add_table(rows=7, cols=3)
    bench_table.style = "Table Grid"
    bench_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    bench_headers = ["Metric", "Ideal Mobile Range", "Ideal Desktop Range"]
    for i, h in enumerate(bench_headers):
        cell = bench_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    bench_data = [
        ["PageSpeed Score", "50-65+", "65-90+"],
        ["CLS", "\u2264 0.1", "\u2264 0.1"],
        ["TBT", "\u2264 800 ms", "\u2264 600 ms"],
        ["FCP", "\u2264 2 sec", "\u2264 1.8 sec"],
        ["LCP", "\u2264 3 sec", "\u2264 2.5 sec"],
        ["Speed Index", "\u2264 4 sec", "\u2264 3 sec"],
    ]
    for r, row_data in enumerate(bench_data):
        for c, val in enumerate(row_data):
            cell = bench_table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ─── Performance Metrics Glossary ───
    doc.add_heading("Performance Metrics Glossary", level=1)
    glossary_table = doc.add_table(rows=6, cols=2)
    glossary_table.style = "Table Grid"
    glossary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    glossary_headers = ["Metric", "Definition"]
    for i, h in enumerate(glossary_headers):
        cell = glossary_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    glossary_data = [
        ["LCP\n(Largest Contentful Paint)", "Measures how long it takes for the main visible content (largest image or text block) to load on the screen."],
        ["FCP\n(First Contentful Paint)", "Indicates when the first visual element appears, showing that the page has started loading."],
        ["TBT\n(Total Blocking Time)", "Calculates how long the page stays unresponsive due to heavy scripts blocking user interaction."],
        ["CLS\n(Cumulative Layout Shift)", "Tracks unexpected layout movements while the page loads, affecting visual stability."],
        ["Speed Index", "Shows how quickly the visible parts of the page are populated during loading."],
    ]
    for r, row_data in enumerate(glossary_data):
        for c, val in enumerate(row_data):
            cell = glossary_table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ─── Expected Performance Outcomes ───
    doc.add_heading("Expected Performance Outcomes", level=1)
    para = doc.add_paragraph()
    run = para.add_run("Based on the current analysis and optimisation scope, the following outcomes are expected:")
    run.font.size = Pt(11)
    para.style = "Body Text"

    outcomes = [
        "Faster and more responsive mobile browsing across Homepage, PDP, and PLP pages",
        "Improved loading priority for key visual elements, allowing important content to appear sooner",
        "Reduced layout shifts, resulting in a more stable and consistent page experience",
        "Smoother interaction readiness with fewer delays caused by heavy scripts and background processes",
        "Better sequencing of scripts and assets, creating a more efficient overall loading flow",
        "Improved Core Web Vitals stability aligned with modern performance benchmarks",
    ]
    for outcome in outcomes:
        doc.add_paragraph(outcome, style="List Bullet")

    para = doc.add_paragraph()
    run = para.add_run(
        "Overall, these improvements are expected to create a smoother and more reliable browsing journey, "
        "helping users explore products with less friction and enhancing their overall experience on the website."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    doc.add_page_break()

    # ─── Timeline ───
    doc.add_heading("Timeline", level=1)
    timeline_table = doc.add_table(rows=5, cols=3)
    timeline_table.style = "Table Grid"
    timeline_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    timeline_headers = ["Phase", "Duration", "Focus"]
    for i, h in enumerate(timeline_headers):
        cell = timeline_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(cell, "D9E2F3")

    timeline_data = [
        ["Phase 1: Cleanup & Script Optimization", "Week 1", "Remove unused apps, defer JS, clean CSS"],
        ["Phase 2: Image & Media Optimization", "Week 2", "Convert \u2192 WEBP, lazy-load, define aspect ratios"],
        ["Phase 3: Server & CDN Setup", "Week 3", "Configure Cloudflare, gzip/Brotli, preconnect"],
        ["Phase 4: QA & Validation", "Week 3", "Lighthouse testing & regression fixes"],
    ]
    for r, row_data in enumerate(timeline_data):
        for c, val in enumerate(row_data):
            cell = timeline_table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_page_break()

    # ─── Deliverables ───
    doc.add_heading("Deliverables", level=1)
    deliverables = [
        "Homepage, PDP & PLP performance optimisation",
        "Script and asset refinement",
        "Media and loading optimisation",
        "Pre- and post-performance validation",
        "Technical optimisation implementation",
        "Technical documentation for maintenance",
    ]
    for d in deliverables:
        doc.add_paragraph(d, style="List Bullet")

    # ─── Assumption ───
    doc.add_heading("Assumption", level=1)
    assumptions = [
        "Optimisation will be carried out within the existing Shopify theme and current website structure.",
        "Third-party apps, marketing scripts, and platform behaviour may influence performance results.",
        "Improvements will aim for the highest feasible optimisation within the current setup.",
        "New features, integrations, or design changes during the optimisation phase may affect performance outcomes.",
        "Core Web Vitals field data take up to 28 days to reflect post-deployment improvements.",
        "Scope is limited to performance optimisation; UI/UX redesign or functional changes are not included unless discussed separately.",
        "Required access to theme files, apps, and tools will be provided for implementation.",
    ]
    for a in assumptions:
        doc.add_paragraph(a, style="List Bullet")

    # ─── Cost Estimate ───
    doc.add_heading("Cost Estimate", level=1)
    cost_table = doc.add_table(rows=1, cols=2)
    cost_table.style = "Table Grid"
    cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cost_table.rows[0].cells[0].text = "Total Cost Estimate"
    cost_table.rows[0].cells[1].text = "INR 1,70,000 + GST"
    for cell in cost_table.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        set_cell_shading(cell, "D9E2F3")

    doc.add_page_break()

    # ─── Conclusion ───
    doc.add_heading("Conclusion", level=1)
    para = doc.add_paragraph()
    run = para.add_run(
        f"{brand_name}'s website demonstrates a strong desktop foundation, with the primary opportunity centred "
        f"around improving mobile performance. Through structured optimisation across scripts, assets, and loading "
        f"behaviour, the website can become more faster with faster load times while maintaining its existing "
        f"design and user experience."
    )
    run.font.size = Pt(11)
    para.style = "Body Text"

    # ─── Logo & Footer ───
    if logo_path:
        add_logo_to_header(doc, logo_path)
    add_footer(doc)

    # ─── Save ───
    if not output_path:
        output_path = f"/Users/dharmikkakadiya/Desktop/{brand_name} Performance Audit Report.docx"

    doc.save(output_path)
    print(f"\nReport saved to: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Shopify Performance Audit Report Generator")
    parser.add_argument("--homepage", required=True, help="Homepage URL")
    parser.add_argument("--pdp", required=True, help="Product Detail Page URL")
    parser.add_argument("--plp", required=True, help="Product Listing Page URL")
    parser.add_argument("--brand", required=True, help="Brand name")
    parser.add_argument("--logo", default=None, help="Path to brand logo image")
    parser.add_argument("--output", default=None, help="Output .docx file path")
    parser.add_argument("--api-key", default=None, help="Google Cloud API key for PageSpeed Insights (overrides built-in key)")

    args = parser.parse_args()

    # Override API key if provided via argument or env variable
    global GOOGLE_API_KEY
    if args.api_key:
        GOOGLE_API_KEY = args.api_key
    elif os.environ.get("GOOGLE_PSI_API_KEY"):
        GOOGLE_API_KEY = os.environ["GOOGLE_PSI_API_KEY"]
    # else: uses the hardcoded default key

    print(f"Shopify Performance Audit Report Generator")
    print(f"==========================================")
    print(f"Brand: {args.brand}")
    print(f"Homepage: {args.homepage}")
    print(f"PDP: {args.pdp}")
    print(f"PLP: {args.plp}")
    print()

    # Run tests
    print("Running Homepage tests...")
    homepage_data = run_tests_for_page(args.homepage, "Homepage")

    print("\nRunning PDP tests...")
    pdp_data = run_tests_for_page(args.pdp, "PDP")

    print("\nRunning PLP tests...")
    plp_data = run_tests_for_page(args.plp, "PLP")

    # Print summary
    print("\n==========================================")
    print("RESULTS SUMMARY")
    print("==========================================")
    for name, data in [("Homepage", homepage_data), ("PDP", pdp_data), ("PLP", plp_data)]:
        m = data["mobile"]
        d = data["desktop"]
        print(f"\n{name}:")
        print(f"  Mobile  - Score: {m['score']}, CLS: {m['cls']}, TBT: {m['tbt']}ms, FCP: {m['fcp']}s, LCP: {m['lcp']}s, SI: {m['si']}s, CWV: {'Pass' if m['cwv_passed'] else 'Fail'}")
        print(f"  Desktop - Score: {d['score']}, CLS: {d['cls']}, TBT: {d['tbt']}ms, FCP: {d['fcp']}s, LCP: {d['lcp']}s, SI: {d['si']}s, CWV: {'Pass' if d['cwv_passed'] else 'Fail'}")

    # Generate report
    print("\nGenerating .docx report...")
    output = generate_report(homepage_data, pdp_data, plp_data, args.brand, args.logo, args.output)

    print(f"\nDone! Report saved at: {output}")


if __name__ == "__main__":
    main()
